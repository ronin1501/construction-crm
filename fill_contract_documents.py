import os
import re

from docx import Document as DocxDocument

from app import create_app
from app.extensions import db
from app.models import Contract, Document as StoredDocument


app = create_app()


def safe_filename_part(value):
    value = str(value or "document")
    value = re.sub(r"[^\wа-яА-ЯёЁ\-]+", "_", value)
    return value.strip("_")[:80]


def add_field(doc, label, value):
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value) if value not in [None, ""] else "—")


def create_contract_docx(contract, filepath):
    doc = DocxDocument()

    doc.add_heading("Карточка договора", level=1)

    add_field(doc, "Номер договора", contract.number)
    add_field(doc, "Наименование договора", contract.title)
    add_field(doc, "Тип договора", contract.contract_type)
    add_field(doc, "Дата договора", contract.contract_date)
    add_field(doc, "Статус", contract.status)

    doc.add_heading("Стороны договора", level=2)
    add_field(doc, "Клиент / заказчик", contract.client.name if contract.client else "—")
    add_field(doc, "Подрядчик / исполнитель", contract.contractor_name)
    add_field(doc, "Поставщик", contract.supplier_name)

    doc.add_heading("Объект и сроки", level=2)
    add_field(doc, "Объект", contract.object.name if contract.object else "—")
    add_field(doc, "Адрес объекта", contract.object.address if contract.object else "—")
    add_field(doc, "Дата начала", contract.start_date)
    add_field(doc, "Дата окончания", contract.end_date)
    add_field(doc, "Место приемки", contract.acceptance_place)

    doc.add_heading("Финансовые условия", level=2)
    add_field(doc, "Сумма договора", contract.amount)
    add_field(doc, "НДС", contract.vat_rate)

    doc.add_heading("Сдача-приемка", level=2)
    add_field(doc, "Условия поставки / выполнения работ", contract.delivery_terms)
    add_field(doc, "Порядок сдачи-приемки", contract.acceptance_procedure)
    add_field(doc, "Гарантийный срок, месяцев", contract.warranty_months)
    add_field(doc, "Санкции / неустойка", contract.penalty_terms)

    doc.add_heading("Описание", level=2)
    doc.add_paragraph(contract.description or "Описание отсутствует.")

    doc.save(filepath)


def create_act_docx(act, filepath):
    contract = act.contract

    doc = DocxDocument()

    doc.add_heading(act.act_type.upper(), level=1)

    add_field(doc, "Номер акта", act.act_number)
    add_field(doc, "Дата акта", act.act_date)
    add_field(doc, "Договор", f"{contract.number} от {contract.contract_date}")
    add_field(doc, "Заказчик", contract.client.name if contract.client else "—")
    add_field(doc, "Подрядчик", contract.contractor_name)
    add_field(doc, "Объект", contract.object.name if contract.object else "—")
    add_field(doc, "Место приемки", act.acceptance_place or contract.acceptance_place)
    add_field(doc, "Наименование работ", act.work_name)

    doc.add_heading("Стоимость работ", level=2)
    add_field(doc, "Стоимость без НДС", act.amount)
    add_field(doc, "Сумма НДС", act.vat_amount)
    add_field(doc, "Итого с НДС", act.total_amount)

    doc.add_heading("Сдача-приемка", level=2)
    add_field(doc, "Статус приемки", act.acceptance_status)
    add_field(doc, "Подписант заказчика", act.signed_by_customer)
    add_field(doc, "Подписант подрядчика", act.signed_by_contractor)
    add_field(doc, "Комментарий", act.comment)

    doc.add_paragraph()
    doc.add_paragraph("Заказчик: ____________________________")
    doc.add_paragraph("Подрядчик: ___________________________")

    doc.save(filepath)


def save_document_record(contract, filename, original_filename, document_type, extracted_text):
    upload_folder = app.config["UPLOAD_FOLDER"]
    filepath = os.path.join(upload_folder, filename)

    existing = StoredDocument.query.filter_by(
        contract_id=contract.id,
        original_filename=original_filename,
        document_type=document_type
    ).first()

    if existing:
        existing.filename = filename
        existing.file_extension = "docx"
        existing.file_size = os.path.getsize(filepath)
        existing.extracted_text = extracted_text
        return existing

    document = StoredDocument(
        filename=filename,
        original_filename=original_filename,
        document_type=document_type,
        file_extension="docx",
        file_size=os.path.getsize(filepath),
        extracted_text=extracted_text,
        contract_id=contract.id
    )

    db.session.add(document)
    return document


with app.app_context():
    upload_folder = app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_folder, exist_ok=True)

    contracts = Contract.query.order_by(Contract.id).all()

    for contract in contracts:
        number_part = safe_filename_part(contract.number)

        contract_filename = f"dogovor_{contract.id}_{number_part}.docx"
        contract_filepath = os.path.join(upload_folder, contract_filename)

        create_contract_docx(contract, contract_filepath)

        save_document_record(
            contract=contract,
            filename=contract_filename,
            original_filename=f"Карточка договора {contract.number}.docx",
            document_type="Договор",
            extracted_text=f"Карточка договора {contract.number}. {contract.title}"
        )

        for act in contract.acts:
            act_part = safe_filename_part(act.act_number)

            act_filename = f"akt_{act.id}_{act_part}.docx"
            act_filepath = os.path.join(upload_folder, act_filename)

            create_act_docx(act, act_filepath)

            save_document_record(
                contract=contract,
                filename=act_filename,
                original_filename=f"{act.act_type} № {act.act_number}.docx",
                document_type=act.act_type,
                extracted_text=f"{act.act_type} № {act.act_number}. {act.work_name or ''}"
            )

    db.session.commit()

    print("Готово! DOCX-документы добавлены ко всем договорам.")