import os
from datetime import date, datetime, timedelta
from decimal import Decimal
from io import BytesIO
from uuid import uuid4

from flask import (
    Blueprint,
    render_template,
    redirect,
    url_for,
    flash,
    request,
    current_app,
    send_from_directory,
    send_file,
)
from flask_login import login_user, logout_user, login_required, current_user
from werkzeug.security import check_password_hash
from werkzeug.utils import secure_filename
from sqlalchemy import or_

from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from app.extensions import db
from app.models import (
    User,
    Client,
    ConstructionObject,
    Contract,
    Document,
    AcceptanceAct,
    Payment,
    Lead,
    Task,
    ActivityLog,
)
from app.forms import (
    LoginForm,
    ClientForm,
    ObjectForm,
    ContractForm,
    DocumentUploadForm,
    AcceptanceActForm,
)

try:
    from app.forms import LeadForm, TaskForm
except ImportError:
    LeadForm = None
    TaskForm = None


main = Blueprint("main", __name__)


# =========================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =========================

def form_value(form, field_name, default=None):
    field = getattr(form, field_name, None)
    return field.data if field is not None else default


def normalize_relation_id(value):
    return value if value and value != 0 else None


def log_activity(action_type, description, client_id=None, contract_id=None, lead_id=None, task_id=None):
    """
    Сохраняет событие в журнал действий.
    Ошибка журнала не должна ломать основную операцию.
    """
    try:
        user_id = current_user.id if current_user and current_user.is_authenticated else None

        log = ActivityLog(
            action_type=action_type,
            description=description,
            user_id=user_id,
            client_id=client_id,
            contract_id=contract_id,
            lead_id=lead_id,
            task_id=task_id,
        )

        db.session.add(log)
        db.session.commit()
    except Exception:
        db.session.rollback()


def set_relation_choices(form):
    """
    Заполняет выпадающие списки связей для заявок и задач.
    """
    clients = [(0, "Не выбран")] + [
        (client.id, client.name)
        for client in Client.query.order_by(Client.name).all()
    ]

    objects = [(0, "Не выбран")] + [
        (obj.id, obj.name)
        for obj in ConstructionObject.query.order_by(ConstructionObject.name).all()
    ]

    contracts = [(0, "Не выбран")] + [
        (contract.id, f"{contract.number} — {contract.title}")
        for contract in Contract.query.order_by(Contract.id.desc()).all()
    ]

    leads = [(0, "Не выбрана")] + [
        (lead.id, f"#{lead.id} — {lead.title}")
        for lead in Lead.query.order_by(Lead.id.desc()).all()
    ]

    if hasattr(form, "client_id"):
        form.client_id.choices = clients

    if hasattr(form, "object_id"):
        form.object_id.choices = objects

    if hasattr(form, "contract_id"):
        form.contract_id.choices = contracts

    if hasattr(form, "lead_id"):
        form.lead_id.choices = leads


def parse_vat_rate(vat_rate):
    if not vat_rate:
        return Decimal("20")

    value = str(vat_rate).replace("%", "").replace(",", ".").strip()

    try:
        return Decimal(value)
    except Exception:
        return Decimal("20")


def calculate_amounts_from_total(total_amount, vat_rate):
    if not total_amount:
        return None, None, None

    total = Decimal(str(total_amount))
    rate = parse_vat_rate(vat_rate)

    if rate <= 0:
        return (
            total.quantize(Decimal("0.01")),
            Decimal("0.00"),
            total.quantize(Decimal("0.01")),
        )

    amount_without_vat = total / (Decimal("1") + rate / Decimal("100"))
    vat_amount = total - amount_without_vat

    return (
        amount_without_vat.quantize(Decimal("0.01")),
        vat_amount.quantize(Decimal("0.01")),
        total.quantize(Decimal("0.01")),
    )


def get_default_act_type(contract):
    if contract.contract_type == "Муниципальный контракт":
        return "КС-2"

    if contract.contract_type == "Договор поставки":
        return "Акт сдачи-приемки"

    return "Акт выполненных работ"


def fill_act_form_by_contract(act_form, contract):
    amount, vat_amount, total_amount = calculate_amounts_from_total(
        contract.amount,
        contract.vat_rate,
    )

    if contract.end_date:
        act_date = contract.end_date
    elif contract.contract_date:
        act_date = contract.contract_date
    else:
        act_date = None

    if contract.status == "Исполнен":
        acceptance_status = "Подписан"
    elif contract.status == "На приемке":
        acceptance_status = "На подписании"
    else:
        acceptance_status = "Подготовлен"

    if contract.contract_type == "Муниципальный контракт":
        work_name = f"Акт о приемке выполненных работ по договору {contract.number}"
        act_number = f"КС-2/{contract.number}"
    elif contract.contract_type == "Договор поставки":
        work_name = f"Поставка материалов по договору {contract.number}"
        act_number = f"УПД/{contract.number}"
    else:
        work_name = contract.title
        act_number = f"Акт/{contract.number}"

    act_form.act_number.data = act_number
    act_form.act_date.data = act_date
    act_form.act_type.data = get_default_act_type(contract)
    act_form.work_name.data = work_name

    act_form.amount.data = amount
    act_form.vat_amount.data = vat_amount
    act_form.total_amount.data = total_amount

    act_form.acceptance_place.data = contract.acceptance_place
    act_form.acceptance_status.data = acceptance_status

    if contract.client and contract.client.representative:
        act_form.signed_by_customer.data = contract.client.representative
    elif contract.client:
        act_form.signed_by_customer.data = contract.client.name
    else:
        act_form.signed_by_customer.data = ""

    act_form.signed_by_contractor.data = "Директор ООО «Дорожник»"

    act_form.comment.data = (
        "Акт сформирован на основании данных договора. "
        "Перед подписанием необходимо проверить объем, стоимость и реквизиты сторон."
    )


def allowed_document(filename):
    if "." not in filename:
        return False

    extension = filename.rsplit(".", 1)[1].lower()

    allowed_extensions = current_app.config.get(
        "ALLOWED_DOCUMENT_EXTENSIONS",
        {"pdf", "docx", "doc", "txt"},
    )

    return extension in allowed_extensions


def extract_text_from_pdf(filepath):
    text = []

    try:
        reader = PdfReader(filepath)

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)

    except Exception:
        return ""

    return "\n".join(text)


def extract_text_from_docx(filepath):
    text = []

    try:
        document = DocxDocument(filepath)

        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text.strip())

    except Exception:
        return ""

    return "\n".join(text)


def extract_text_from_txt(filepath):
    try:
        with open(filepath, "r", encoding="utf-8") as file:
            return file.read()
    except Exception:
        return ""


def extract_text_from_file(filepath, extension):
    if extension == "pdf":
        return extract_text_from_pdf(filepath)

    if extension == "docx":
        return extract_text_from_docx(filepath)

    if extension == "txt":
        return extract_text_from_txt(filepath)

    return ""


def save_document_file(file, contract_id, document_type):
    original_filename = file.filename
    safe_name = secure_filename(original_filename)

    if not safe_name:
        safe_name = f"document_{uuid4().hex}"

    extension = safe_name.rsplit(".", 1)[1].lower()
    stored_filename = f"{uuid4().hex}_{safe_name}"

    upload_folder = current_app.config["UPLOAD_FOLDER"]
    os.makedirs(upload_folder, exist_ok=True)

    filepath = os.path.join(upload_folder, stored_filename)
    file.save(filepath)

    file_size = os.path.getsize(filepath)
    extracted_text = extract_text_from_file(filepath, extension)

    document = Document(
        filename=stored_filename,
        original_filename=original_filename,
        document_type=document_type,
        file_extension=extension,
        file_size=file_size,
        extracted_text=extracted_text,
        contract_id=contract_id,
    )

    db.session.add(document)
    db.session.commit()

    log_activity(
        "Загрузка документа",
        f"Загружен документ «{original_filename}» по договору ID {contract_id}.",
        contract_id=contract_id,
    )

    return document


def add_docx_field(document, label, value):
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value) if value not in [None, ""] else "—")


def make_docx_response(document, filename):
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def calculate_dashboard_metrics():
    today = date.today()
    next_30_days = today + timedelta(days=30)

    expected_payments_total = db.session.query(db.func.sum(Payment.amount)).filter(
        Payment.status.in_(["Ожидается", "Просрочено"])
    ).scalar() or 0

    paid_payments_total = db.session.query(db.func.sum(Payment.amount)).filter(
        Payment.status == "Оплачено"
    ).scalar() or 0

    overdue_payments_count = Payment.query.filter(
        Payment.status == "Ожидается",
        Payment.payment_date.isnot(None),
        Payment.payment_date < today,
    ).count()

    expiring_contracts_count = Contract.query.filter(
        Contract.end_date.isnot(None),
        Contract.end_date >= today,
        Contract.end_date <= next_30_days,
        Contract.status.notin_(["Исполнен", "Расторгнут"]),
    ).count()

    return {
        "clients_count": Client.query.count(),
        "contracts_count": Contract.query.count(),
        "objects_count": ConstructionObject.query.count(),
        "active_contracts": Contract.query.filter(
            Contract.status.in_(["Действующий", "На исполнении", "На приемке"])
        ).count(),
        "new_leads_count": Lead.query.filter_by(status="Новая").count(),
        "tasks_today_count": Task.query.filter(
            Task.due_date == today,
            Task.status != "Выполнена",
        ).count(),
        "overdue_tasks_count": Task.query.filter(
            Task.due_date.isnot(None),
            Task.due_date < today,
            Task.status.notin_(["Выполнена", "Отменена"]),
        ).count(),
        "acts_on_signing_count": AcceptanceAct.query.filter_by(
            acceptance_status="На подписании"
        ).count(),
        "expected_payments_total": expected_payments_total,
        "paid_payments_total": paid_payments_total,
        "overdue_payments_count": overdue_payments_count,
        "expiring_contracts_count": expiring_contracts_count,
    }


# =========================
# ГЛАВНАЯ И АВТОРИЗАЦИЯ
# =========================

@main.route("/")
@login_required
def index():
    metrics = calculate_dashboard_metrics()

    recent_leads = Lead.query.order_by(Lead.id.desc()).limit(5).all()

    upcoming_tasks = Task.query.filter(
        Task.status.notin_(["Выполнена", "Отменена"])
    ).order_by(
        Task.due_date.asc().nullslast(),
        Task.id.desc(),
    ).limit(5).all()

    return render_template(
        "index.html",
        metrics=metrics,
        recent_leads=recent_leads,
        upcoming_tasks=upcoming_tasks,
    )


@main.route("/login", methods=["GET", "POST"])
def login():
    form = LoginForm()

    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()

        if user and check_password_hash(user.password_hash, form.password.data):
            login_user(user)
            return redirect(url_for("main.index"))

        flash("Неверный логин или пароль", "danger")

    return render_template("login.html", form=form)


@main.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("main.login"))


# =========================
# КЛИЕНТЫ
# =========================

@main.route("/clients")
@login_required
def clients_list():
    search = request.args.get("search", "").strip()
    tag = request.args.get("tag", "").strip()

    query = Client.query

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Client.name.ilike(pattern),
                Client.inn.ilike(pattern),
                Client.phone.ilike(pattern),
                Client.email.ilike(pattern),
                Client.tags_text.ilike(pattern),
            )
        )

    if tag:
        query = query.filter(Client.tags_text.ilike(f"%{tag}%"))

    clients = query.order_by(Client.id.desc()).all()

    return render_template(
        "clients/clients_list.html",
        clients=clients,
        search=search,
        tag=tag,
    )


@main.route("/clients/add", methods=["GET", "POST"])
@login_required
def client_add():
    form = ClientForm()

    if form.validate_on_submit():
        client = Client(
            client_type=form.client_type.data,
            name=form.name.data,
            inn=form_value(form, "inn"),
            kpp=form_value(form, "kpp"),
            ogrn=form_value(form, "ogrn"),
            phone=form_value(form, "phone"),
            email=form_value(form, "email"),
            address=form_value(form, "address"),
            representative=form_value(form, "representative"),
            tags_text=form_value(form, "tags_text"),
            note=form_value(form, "note"),
        )

        db.session.add(client)
        db.session.commit()

        log_activity(
            "Создание клиента",
            f"Добавлен клиент «{client.name}».",
            client_id=client.id,
        )

        flash("Клиент добавлен", "success")
        return redirect(url_for("main.clients_list"))

    return render_template("clients/client_add.html", form=form)


@main.route("/clients/<int:id>")
@login_required
def client_detail(id):
    client = Client.query.get_or_404(id)
    return render_template("clients/client_detail.html", client=client)


@main.route("/clients/<int:id>/edit", methods=["GET", "POST"])
@login_required
def client_edit(id):
    client = Client.query.get_or_404(id)
    form = ClientForm(obj=client)

    if form.validate_on_submit():
        form.populate_obj(client)
        db.session.commit()

        log_activity(
            "Изменение клиента",
            f"Обновлены данные клиента «{client.name}».",
            client_id=client.id,
        )

        flash("Данные клиента обновлены", "success")
        return redirect(url_for("main.client_detail", id=client.id))

    return render_template("clients/client_edit.html", form=form, client=client)


# =========================
# ЗАЯВКИ / ОБРАЩЕНИЯ
# =========================

@main.route("/leads")
@login_required
def leads_list():
    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()
    source = request.args.get("source", "").strip()
    priority = request.args.get("priority", "").strip()

    query = Lead.query.outerjoin(Client)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Lead.title.ilike(pattern),
                Lead.contact_person.ilike(pattern),
                Lead.phone.ilike(pattern),
                Lead.email.ilike(pattern),
                Lead.work_type.ilike(pattern),
                Lead.object_address.ilike(pattern),
                Client.name.ilike(pattern),
            )
        )

    if status:
        query = query.filter(Lead.status == status)

    if source:
        query = query.filter(Lead.source == source)

    if priority:
        query = query.filter(Lead.priority == priority)

    leads = query.order_by(Lead.lead_date.desc(), Lead.id.desc()).all()

    return render_template(
        "leads/leads_list.html",
        leads=leads,
        search=search,
        status=status,
        source=source,
        priority=priority,
    )


@main.route("/leads/add", methods=["GET", "POST"])
@login_required
def lead_add():
    if LeadForm is None:
        flash("Форма заявок пока не подключена. Сначала обнови app/forms.py.", "danger")
        return redirect(url_for("main.leads_list"))

    form = LeadForm()
    set_relation_choices(form)

    if form.validate_on_submit():
        lead = Lead(
            title=form.title.data,
            lead_date=form.lead_date.data,
            source=form.source.data,
            status=form.status.data,
            priority=form.priority.data,
            contact_person=form.contact_person.data,
            phone=form.phone.data,
            email=form.email.data,
            work_type=form.work_type.data,
            object_address=form.object_address.data,
            estimated_budget=form.estimated_budget.data,
            responsible=form.responsible.data,
            comment=form.comment.data,
            client_id=normalize_relation_id(form.client_id.data),
            object_id=normalize_relation_id(form.object_id.data),
            contract_id=normalize_relation_id(form.contract_id.data),
        )

        db.session.add(lead)
        db.session.commit()

        log_activity(
            "Создание заявки",
            f"Создана заявка #{lead.id}: {lead.title}.",
            client_id=lead.client_id,
            contract_id=lead.contract_id,
            lead_id=lead.id,
        )

        flash("Заявка добавлена", "success")
        return redirect(url_for("main.lead_detail", id=lead.id))

    return render_template("leads/lead_add.html", form=form)


@main.route("/leads/<int:id>")
@login_required
def lead_detail(id):
    lead = Lead.query.get_or_404(id)
    return render_template("leads/lead_detail.html", lead=lead)


@main.route("/leads/<int:id>/edit", methods=["GET", "POST"])
@login_required
def lead_edit(id):
    if LeadForm is None:
        flash("Форма заявок пока не подключена. Сначала обнови app/forms.py.", "danger")
        return redirect(url_for("main.lead_detail", id=id))

    lead = Lead.query.get_or_404(id)
    form = LeadForm(obj=lead)
    set_relation_choices(form)

    if request.method == "GET":
        form.client_id.data = lead.client_id or 0
        form.object_id.data = lead.object_id or 0
        form.contract_id.data = lead.contract_id or 0

    if form.validate_on_submit():
        lead.title = form.title.data
        lead.lead_date = form.lead_date.data
        lead.source = form.source.data
        lead.status = form.status.data
        lead.priority = form.priority.data
        lead.contact_person = form.contact_person.data
        lead.phone = form.phone.data
        lead.email = form.email.data
        lead.work_type = form.work_type.data
        lead.object_address = form.object_address.data
        lead.estimated_budget = form.estimated_budget.data
        lead.responsible = form.responsible.data
        lead.comment = form.comment.data
        lead.client_id = normalize_relation_id(form.client_id.data)
        lead.object_id = normalize_relation_id(form.object_id.data)
        lead.contract_id = normalize_relation_id(form.contract_id.data)

        db.session.commit()

        log_activity(
            "Изменение заявки",
            f"Обновлена заявка #{lead.id}: статус «{lead.status}».",
            client_id=lead.client_id,
            contract_id=lead.contract_id,
            lead_id=lead.id,
        )

        flash("Заявка обновлена", "success")
        return redirect(url_for("main.lead_detail", id=lead.id))

    return render_template("leads/lead_edit.html", form=form, lead=lead)


@main.route("/leads/<int:id>/delete", methods=["POST"])
@login_required
def lead_delete(id):
    lead = Lead.query.get_or_404(id)

    lead_title = lead.title

    Task.query.filter_by(lead_id=lead.id).update({"lead_id": None})
    ActivityLog.query.filter_by(lead_id=lead.id).update({"lead_id": None})

    db.session.delete(lead)
    db.session.commit()

    log_activity(
        "Удаление заявки",
        f"Удалена заявка: {lead_title}.",
    )

    flash("Заявка удалена", "success")
    return redirect(url_for("main.leads_list"))


@main.route("/leads/<int:id>/task", methods=["POST"])
@login_required
def lead_create_task(id):
    lead = Lead.query.get_or_404(id)

    task = Task(
        title=f"Обработать заявку: {lead.title}",
        description=(
            f"Связаться с клиентом по заявке.\n\n"
            f"Контактное лицо: {lead.contact_person or '—'}\n"
            f"Телефон: {lead.phone or '—'}\n"
            f"Email: {lead.email or '—'}\n"
            f"Вид работ: {lead.work_type or '—'}\n"
            f"Комментарий: {lead.comment or '—'}"
        ),
        task_type="Звонок клиенту",
        priority=lead.priority or "Обычная",
        status="Новая",
        due_date=date.today() + timedelta(days=2),
        responsible=lead.responsible or "Менеджер",
        client_id=lead.client_id,
        contract_id=lead.contract_id,
        lead_id=lead.id,
    )

    db.session.add(task)
    db.session.commit()

    log_activity(
        "Создание задачи по заявке",
        f"По заявке #{lead.id} создана задача: {task.title}.",
        client_id=task.client_id,
        contract_id=task.contract_id,
        lead_id=lead.id,
        task_id=task.id,
    )

    flash("Задача по заявке создана", "success")
    return redirect(url_for("main.lead_detail", id=lead.id))


# =========================
# ЗАДАЧИ И НАПОМИНАНИЯ
# =========================

@main.route("/tasks")
@login_required
def tasks_list():
    today = date.today()

    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()
    priority = request.args.get("priority", "").strip()

    query = Task.query.outerjoin(Task.client).outerjoin(Task.contract).outerjoin(Task.lead)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Task.title.ilike(pattern),
                Task.description.ilike(pattern),
                Task.responsible.ilike(pattern),
                Client.name.ilike(pattern),
                Contract.number.ilike(pattern),
                Lead.title.ilike(pattern),
                Lead.contact_person.ilike(pattern),
            )
        )

    if status:
        query = query.filter(Task.status == status)

    if priority:
        query = query.filter(Task.priority == priority)

    tasks = query.order_by(
        Task.due_date.asc().nullslast(),
        Task.id.desc(),
    ).all()

    return render_template(
        "tasks/tasks_list.html",
        tasks=tasks,
        search=search,
        status=status,
        priority=priority,
        today=today,
    )


@main.route("/tasks/add", methods=["GET", "POST"])
@login_required
def task_add():
    if TaskForm is None:
        flash("Форма задач пока не подключена. Сначала обнови app/forms.py.", "danger")
        return redirect(url_for("main.tasks_list"))

    form = TaskForm()
    set_relation_choices(form)

    if form.validate_on_submit():
        task = Task(
            title=form.title.data,
            description=form.description.data,
            task_type=form.task_type.data,
            priority=form.priority.data,
            status=form.status.data,
            due_date=form.due_date.data,
            responsible=form.responsible.data,
            client_id=normalize_relation_id(form.client_id.data),
            contract_id=normalize_relation_id(form.contract_id.data),
            lead_id=normalize_relation_id(form.lead_id.data),
        )

        db.session.add(task)
        db.session.commit()

        log_activity(
            "Создание задачи",
            f"Создана задача #{task.id}: {task.title}.",
            client_id=task.client_id,
            contract_id=task.contract_id,
            lead_id=task.lead_id,
            task_id=task.id,
        )

        flash("Задача добавлена", "success")
        return redirect(url_for("main.tasks_list"))

    return render_template("tasks/task_add.html", form=form)


@main.route("/tasks/<int:id>/edit", methods=["GET", "POST"])
@login_required
def task_edit(id):
    if TaskForm is None:
        flash("Форма задач пока не подключена. Сначала обнови app/forms.py.", "danger")
        return redirect(url_for("main.tasks_list"))

    task = Task.query.get_or_404(id)
    form = TaskForm(obj=task)
    set_relation_choices(form)

    if request.method == "GET":
        form.client_id.data = task.client_id or 0
        form.contract_id.data = task.contract_id or 0
        form.lead_id.data = task.lead_id or 0

    if form.validate_on_submit():
        task.title = form.title.data
        task.description = form.description.data
        task.task_type = form.task_type.data
        task.priority = form.priority.data
        task.status = form.status.data
        task.due_date = form.due_date.data
        task.responsible = form.responsible.data
        task.client_id = normalize_relation_id(form.client_id.data)
        task.contract_id = normalize_relation_id(form.contract_id.data)
        task.lead_id = normalize_relation_id(form.lead_id.data)

        if task.status == "Выполнена" and not task.completed_at:
            task.completed_at = datetime.utcnow()

        if task.status != "Выполнена":
            task.completed_at = None

        db.session.commit()

        log_activity(
            "Изменение задачи",
            f"Обновлена задача #{task.id}: статус «{task.status}».",
            client_id=task.client_id,
            contract_id=task.contract_id,
            lead_id=task.lead_id,
            task_id=task.id,
        )

        flash("Задача обновлена", "success")
        return redirect(url_for("main.tasks_list"))

    return render_template("tasks/task_edit.html", form=form, task=task)


@main.route("/tasks/<int:id>/complete", methods=["POST", "GET"])
@login_required
def task_complete(id):
    task = Task.query.get_or_404(id)

    task.status = "Выполнена"
    task.completed_at = datetime.utcnow()

    db.session.commit()

    log_activity(
        "Выполнение задачи",
        f"Задача #{task.id} отмечена как выполненная.",
        client_id=task.client_id,
        contract_id=task.contract_id,
        lead_id=task.lead_id,
        task_id=task.id,
    )

    flash("Задача отмечена как выполненная", "success")
    return redirect(url_for("main.tasks_list"))


# =========================
# ДОГОВОРЫ
# =========================

@main.route("/contracts")
@login_required
def contracts_list():
    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()

    query = Contract.query.join(Client)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Contract.number.ilike(pattern),
                Contract.title.ilike(pattern),
                Contract.contract_type.ilike(pattern),
                Client.name.ilike(pattern),
                Client.inn.ilike(pattern),
            )
        )

    if status:
        query = query.filter(Contract.status == status)

    contracts = query.order_by(Contract.id.desc()).all()

    return render_template(
        "contracts/contracts_list.html",
        contracts=contracts,
        search=search,
        status=status,
    )


@main.route("/contracts/add", methods=["GET", "POST"])
@login_required
def contract_add():
    form = ContractForm()

    form.client_id.choices = [
        (client.id, client.name)
        for client in Client.query.order_by(Client.name).all()
    ]

    form.object_id.choices = [(0, "Без объекта")] + [
        (obj.id, obj.name)
        for obj in ConstructionObject.query.order_by(ConstructionObject.name).all()
    ]

    if form.validate_on_submit():
        contract = Contract(
            number=form.number.data,
            title=form.title.data,
            contract_type=form.contract_type.data,
            contract_date=form.contract_date.data,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            amount=form.amount.data,
            vat_rate=form.vat_rate.data,
            status=form.status.data,
            client_id=form.client_id.data,
            object_id=form.object_id.data if form.object_id.data != 0 else None,
            contractor_name=form.contractor_name.data,
            supplier_name=form.supplier_name.data,
            acceptance_place=form.acceptance_place.data,
            delivery_terms=form.delivery_terms.data,
            acceptance_procedure=form.acceptance_procedure.data,
            warranty_months=form.warranty_months.data,
            penalty_terms=form.penalty_terms.data,
            description=form.description.data,
        )

        db.session.add(contract)
        db.session.commit()

        if form.document.data:
            file = form.document.data

            if file.filename and allowed_document(file.filename):
                save_document_file(
                    file=file,
                    contract_id=contract.id,
                    document_type="Договор",
                )

        log_activity(
            "Создание договора",
            f"Добавлен договор №{contract.number}: {contract.title}.",
            client_id=contract.client_id,
            contract_id=contract.id,
        )

        flash("Договор добавлен", "success")
        return redirect(url_for("main.contracts_list"))

    return render_template("contracts/contract_add.html", form=form)


@main.route("/contracts/<int:id>")
@login_required
def contract_detail(id):
    contract = Contract.query.get_or_404(id)

    upload_form = DocumentUploadForm()
    act_form = AcceptanceActForm()

    fill_act_form_by_contract(act_form, contract)

    return render_template(
        "contracts/contract_detail.html",
        contract=contract,
        upload_form=upload_form,
        act_form=act_form,
    )


@main.route("/contracts/<int:id>/edit", methods=["GET", "POST"])
@login_required
def contract_edit(id):
    contract = Contract.query.get_or_404(id)
    form = ContractForm(obj=contract)

    form.client_id.choices = [
        (client.id, client.name)
        for client in Client.query.order_by(Client.name).all()
    ]

    form.object_id.choices = [(0, "Без объекта")] + [
        (obj.id, obj.name)
        for obj in ConstructionObject.query.order_by(ConstructionObject.name).all()
    ]

    if request.method == "GET":
        form.client_id.data = contract.client_id
        form.object_id.data = contract.object_id or 0

    if form.validate_on_submit():
        contract.number = form.number.data
        contract.title = form.title.data
        contract.contract_type = form.contract_type.data
        contract.contract_date = form.contract_date.data
        contract.start_date = form.start_date.data
        contract.end_date = form.end_date.data
        contract.amount = form.amount.data
        contract.vat_rate = form.vat_rate.data
        contract.status = form.status.data
        contract.client_id = form.client_id.data
        contract.object_id = form.object_id.data if form.object_id.data != 0 else None
        contract.contractor_name = form.contractor_name.data
        contract.supplier_name = form.supplier_name.data
        contract.acceptance_place = form.acceptance_place.data
        contract.delivery_terms = form.delivery_terms.data
        contract.acceptance_procedure = form.acceptance_procedure.data
        contract.warranty_months = form.warranty_months.data
        contract.penalty_terms = form.penalty_terms.data
        contract.description = form.description.data

        db.session.commit()

        log_activity(
            "Изменение договора",
            f"Обновлен договор №{contract.number}.",
            client_id=contract.client_id,
            contract_id=contract.id,
        )

        flash("Договор обновлен", "success")
        return redirect(url_for("main.contract_detail", id=contract.id))

    return render_template("contracts/contract_edit.html", form=form, contract=contract)


@main.route("/contracts/<int:id>/export/docx")
@login_required
def contract_export_docx(id):
    contract = Contract.query.get_or_404(id)

    document = DocxDocument()
    document.add_heading("Карточка договора", level=1)

    add_docx_field(document, "Номер договора", contract.number)
    add_docx_field(document, "Наименование", contract.title)
    add_docx_field(document, "Тип договора", contract.contract_type)
    add_docx_field(document, "Дата договора", contract.contract_date)
    add_docx_field(document, "Статус", contract.status)

    document.add_heading("Стороны договора", level=2)
    add_docx_field(document, "Заказчик / клиент", contract.client.name if contract.client else "")
    add_docx_field(document, "Подрядчик / исполнитель", contract.contractor_name)
    add_docx_field(document, "Поставщик", contract.supplier_name)

    document.add_heading("Объект и сроки", level=2)
    add_docx_field(document, "Объект", contract.object.name if contract.object else "")
    add_docx_field(document, "Адрес объекта", contract.object.address if contract.object else "")
    add_docx_field(document, "Дата начала", contract.start_date)
    add_docx_field(document, "Дата окончания", contract.end_date)
    add_docx_field(document, "Место приемки", contract.acceptance_place)

    document.add_heading("Финансовые условия", level=2)
    add_docx_field(document, "Сумма договора", contract.amount)
    add_docx_field(document, "НДС", contract.vat_rate)

    document.add_heading("Условия исполнения", level=2)
    add_docx_field(document, "Условия поставки / выполнения работ", contract.delivery_terms)
    add_docx_field(document, "Порядок сдачи-приемки", contract.acceptance_procedure)
    add_docx_field(document, "Гарантийный срок, месяцев", contract.warranty_months)
    add_docx_field(document, "Санкции / неустойка", contract.penalty_terms)
    add_docx_field(document, "Описание", contract.description)

    document.add_heading("Документы по договору", level=2)

    if contract.documents:
        for doc in contract.documents:
            add_docx_field(
                document,
                doc.document_type or "Документ",
                doc.original_filename or doc.filename,
            )
    else:
        document.add_paragraph("Документы не прикреплены.")

    safe_number = str(contract.number).replace("/", "-").replace("\\", "-")
    filename = f"dogovor_{safe_number}.docx"

    return make_docx_response(document, filename)


# =========================
# ДОКУМЕНТЫ
# =========================

@main.route("/contracts/<int:id>/documents/add", methods=["POST"])
@login_required
def document_add(id):
    contract = Contract.query.get_or_404(id)
    form = DocumentUploadForm()

    if form.validate_on_submit():
        file = form.file.data

        if not file or file.filename == "":
            flash("Файл не выбран", "danger")
            return redirect(url_for("main.contract_detail", id=contract.id))

        if not allowed_document(file.filename):
            flash("Можно загружать только файлы PDF, DOCX, DOC или TXT", "danger")
            return redirect(url_for("main.contract_detail", id=contract.id))

        save_document_file(
            file=file,
            contract_id=contract.id,
            document_type=form.document_type.data,
        )

        flash("Документ успешно загружен", "success")
        return redirect(url_for("main.contract_detail", id=contract.id))

    flash("Ошибка при загрузке документа", "danger")
    return redirect(url_for("main.contract_detail", id=contract.id))


@main.route("/documents/<int:id>/download")
@login_required
def document_download(id):
    document = Document.query.get_or_404(id)

    upload_folder = current_app.config["UPLOAD_FOLDER"]
    file_path = os.path.join(upload_folder, document.filename)

    if not os.path.exists(file_path):
        flash("Файл документа не найден в папке uploads/documents", "danger")
        return redirect(url_for("main.contract_detail", id=document.contract_id))

    return send_from_directory(
        upload_folder,
        document.filename,
        as_attachment=True,
        download_name=document.original_filename or document.filename,
    )


@main.route("/documents")
@login_required
def documents_list():
    search = request.args.get("search", "").strip()
    document_type = request.args.get("document_type", "").strip()

    query = Document.query.join(Contract)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Document.original_filename.ilike(pattern),
                Document.filename.ilike(pattern),
                Contract.number.ilike(pattern),
                Contract.title.ilike(pattern),
            )
        )

    if document_type:
        query = query.filter(Document.document_type == document_type)

    documents = query.order_by(Document.id.desc()).all()

    return render_template(
        "documents/documents_list.html",
        documents=documents,
        search=search,
        document_type=document_type,
    )


# =========================
# АКТЫ
# =========================

@main.route("/contracts/<int:id>/acts/add", methods=["POST"])
@login_required
def act_add(id):
    contract = Contract.query.get_or_404(id)
    form = AcceptanceActForm()

    if form.validate_on_submit():
        act = AcceptanceAct(
            act_number=form.act_number.data,
            act_date=form.act_date.data,
            act_type=form.act_type.data,
            work_name=form.work_name.data,
            amount=form.amount.data,
            vat_amount=form.vat_amount.data,
            total_amount=form.total_amount.data,
            acceptance_place=form.acceptance_place.data,
            acceptance_status=form.acceptance_status.data,
            signed_by_customer=form.signed_by_customer.data,
            signed_by_contractor=form.signed_by_contractor.data,
            comment=form.comment.data,
            contract_id=contract.id,
        )

        db.session.add(act)
        db.session.commit()

        log_activity(
            "Создание акта",
            f"Добавлен {act.act_type} №{act.act_number}.",
            client_id=contract.client_id,
            contract_id=contract.id,
        )

        flash("Акт выполненных работ добавлен", "success")
        return redirect(url_for("main.contract_detail", id=contract.id))

    flash("Ошибка при добавлении акта", "danger")
    return redirect(url_for("main.contract_detail", id=contract.id))


@main.route("/acts/<int:id>/export/docx")
@login_required
def act_export_docx(id):
    act = AcceptanceAct.query.get_or_404(id)
    contract = act.contract

    document = DocxDocument()
    document.add_heading(act.act_type.upper(), level=1)

    add_docx_field(document, "Номер акта", act.act_number)
    add_docx_field(document, "Дата акта", act.act_date)
    add_docx_field(document, "Договор", f"{contract.number} от {contract.contract_date}")
    add_docx_field(document, "Заказчик", contract.client.name if contract.client else "")
    add_docx_field(document, "Подрядчик", contract.contractor_name)
    add_docx_field(document, "Объект", contract.object.name if contract.object else "")
    add_docx_field(document, "Место приемки", act.acceptance_place or contract.acceptance_place)
    add_docx_field(document, "Наименование работ", act.work_name)

    document.add_heading("Стоимость работ", level=2)
    add_docx_field(document, "Стоимость без НДС", act.amount)
    add_docx_field(document, "НДС", act.vat_amount)
    add_docx_field(document, "Итого с НДС", act.total_amount)

    document.add_heading("Сдача-приемка", level=2)
    add_docx_field(document, "Статус приемки", act.acceptance_status)
    add_docx_field(document, "Подписант заказчика", act.signed_by_customer)
    add_docx_field(document, "Подписант подрядчика", act.signed_by_contractor)
    add_docx_field(document, "Комментарий", act.comment)

    document.add_paragraph()
    document.add_paragraph("Заказчик: ____________________________")
    document.add_paragraph("Подрядчик: ___________________________")

    safe_number = str(act.act_number).replace("/", "-").replace("\\", "-")
    filename = f"akt_{safe_number}.docx"

    return make_docx_response(document, filename)


@main.route("/acts")
@login_required
def acts_list():
    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()

    query = AcceptanceAct.query.join(Contract)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                AcceptanceAct.act_number.ilike(pattern),
                AcceptanceAct.work_name.ilike(pattern),
                Contract.number.ilike(pattern),
                Contract.title.ilike(pattern),
            )
        )

    if status:
        query = query.filter(AcceptanceAct.acceptance_status == status)

    acts = query.order_by(AcceptanceAct.id.desc()).all()

    return render_template(
        "acts/acts_list.html",
        acts=acts,
        search=search,
        status=status,
    )


# =========================
# ОБЪЕКТЫ
# =========================

@main.route("/objects")
@login_required
def objects_list():
    objects = ConstructionObject.query.order_by(ConstructionObject.id.desc()).all()
    return render_template("objects/objects_list.html", objects=objects)


@main.route("/objects/add", methods=["GET", "POST"])
@login_required
def object_add():
    form = ObjectForm()

    if form.validate_on_submit():
        obj = ConstructionObject(
            name=form.name.data,
            address=form_value(form, "address"),
            district=form_value(form, "district"),
            work_type=form_value(form, "work_type"),
            status=form.status.data,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
        )

        db.session.add(obj)
        db.session.commit()

        log_activity("Создание объекта", f"Добавлен объект «{obj.name}».")

        flash("Объект добавлен", "success")
        return redirect(url_for("main.objects_list"))

    return render_template("objects/object_add.html", form=form)


@main.route("/objects/<int:id>")
@login_required
def object_detail(id):
    obj = ConstructionObject.query.get_or_404(id)
    return render_template("objects/object_detail.html", obj=obj)


@main.route("/objects/<int:id>/edit", methods=["GET", "POST"])
@login_required
def object_edit(id):
    obj = ConstructionObject.query.get_or_404(id)
    form = ObjectForm(obj=obj)

    if form.validate_on_submit():
        form.populate_obj(obj)
        db.session.commit()

        log_activity("Изменение объекта", f"Обновлен объект «{obj.name}».")

        flash("Объект обновлен", "success")
        return redirect(url_for("main.object_detail", id=obj.id))

    return render_template("objects/object_edit.html", form=form, obj=obj)


# =========================
# ОПЛАТЫ, ОТЧЕТЫ, ЖУРНАЛ
# =========================

@main.route("/payments")
@login_required
def payments_list():
    today = date.today()

    search = request.args.get("search", "").strip()
    status = request.args.get("status", "").strip()

    query = Payment.query.join(Contract).join(Client)

    if search:
        pattern = f"%{search}%"
        query = query.filter(
            or_(
                Payment.purpose.ilike(pattern),
                Payment.payment_type.ilike(pattern),
                Contract.number.ilike(pattern),
                Contract.title.ilike(pattern),
                Client.name.ilike(pattern),
            )
        )

    if status:
        query = query.filter(Payment.status == status)

    payments = query.order_by(
        Payment.payment_date.desc().nullslast(),
        Payment.id.desc(),
    ).all()

    total_amount = sum((payment.amount or Decimal("0")) for payment in payments)
    expected_amount = sum(
        (payment.amount or Decimal("0"))
        for payment in payments
        if payment.status == "Ожидается"
    )
    paid_amount = sum(
        (payment.amount or Decimal("0"))
        for payment in payments
        if payment.status == "Оплачено"
    )

    return render_template(
        "payments/payments_list.html",
        payments=payments,
        search=search,
        status=status,
        today=today,
        total_amount=total_amount,
        expected_amount=expected_amount,
        paid_amount=paid_amount,
    )


@main.route("/activity")
@login_required
def activity_list():
    action_type = request.args.get("action_type", "").strip()

    query = ActivityLog.query

    if action_type:
        query = query.filter(ActivityLog.action_type == action_type)

    logs = query.order_by(ActivityLog.id.desc()).limit(300).all()

    return render_template(
        "activity/activity_list.html",
        logs=logs,
        action_type=action_type,
    )


@main.route("/reports")
@login_required
def reports():
    metrics = calculate_dashboard_metrics()
    total_amount = db.session.query(db.func.sum(Contract.amount)).scalar() or 0

    contracts_by_status = db.session.query(
        Contract.status,
        db.func.count(Contract.id),
    ).group_by(Contract.status).all()

    payments_by_status = db.session.query(
        Payment.status,
        db.func.count(Payment.id),
        db.func.coalesce(db.func.sum(Payment.amount), 0),
    ).group_by(Payment.status).all()

    return render_template(
        "reports/reports.html",
        clients_count=Client.query.count(),
        contracts_count=Contract.query.count(),
        objects_count=ConstructionObject.query.count(),
        total_amount=total_amount,
        metrics=metrics,
        contracts_by_status=contracts_by_status,
        payments_by_status=payments_by_status,
    )